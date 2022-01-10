from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from pandas import DataFrame


def write_data(doc, item, view_count=None, error_count=None, high_light=False, show_count=True, show_error=True,
               show_rate=True):
    doc.add_paragraph(item['context'])

    p = doc.add_paragraph()
    text_count = len(''.join(item['option']))
    for i, j in enumerate(item['option']):
        run = p.add_run(f"({chr(i + ord('A'))}) {j}")
        if high_light:
            if i in item["answer"]:
                font = run.font
                font.color.rgb = RGBColor(0xe4, 0x00, 0x2b)
        if i + 1 < len(item['option']):
            if text_count > 25:
                p.add_run('\n')
            else:
                p.add_run('\t')

    p = doc.add_paragraph('答案：' + ''.join(chr(i + ord('A')) for i in item["answer"]))
    if error_count is not None and show_error:
        p.add_run(f' 错误次数：{error_count}')
    if view_count is not None and show_count:
        p.add_run(f' 答题次数：{view_count}')
    if view_count is not None and error_count is not None and show_rate:
        p.add_run(f' 错误率：{error_count / (view_count + 1e-8) * 100:.2f}%')
    if 'knowledge' in item.keys():
        p.add_run(' 知识点：' + item['knowledge'])


def export_data(data, foot: DataFrame, save_path, count=10):
    doc = Document()

    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for index, item in foot.sort_values('error', ascending=False).iloc[:count].iterrows():
        write_data(doc, data[index], item['view'], item['error'], high_light=True)
        doc.add_paragraph()
    doc.save(save_path)
